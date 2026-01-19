import os
import docx
import torch
from transformers import AutoModelForTokenClassification, AutoTokenizer
import re
from datetime import datetime
import sqlite3
from cryptography.fernet import Fernet
import secrets
import string
from typing import List, Tuple
import uuid
from dataclasses import dataclass

INPUT_FOLDER = "input_docs"
OUTPUT_FOLDER = "output_docs"
MODEL_PATH = "models/court_ner_model"
DATABASE_FILE = "anonymization_mapping.db"

@dataclass
class EntityMapping:
    original_text: str
    anonymized_text: str
    entity_type: str
    document_id: str
    position: Tuple[int, int]
    timestamp: datetime

class DataGenerator:
    def __init__(self):
        self.russian_male_names = ['Александр', 'Алексей', 'Андрей', 'Антон', 'Артём', 'Борис']
        self.russian_female_names = ['Александра', 'Алина', 'Алиса', 'Алла', 'Анастасия', 'Анна']
        self.russian_surnames = ['Иванов', 'Петров', 'Сидоров', 'Смирнов', 'Кузнецов', 'Попов']
        self.cities = ['Москва', 'Санкт-Петербург', 'Новосибирск', 'Екатеринбург']
        self.streets = ['Ленина', 'Советская', 'Мира', 'Центральная']

    def generate_fio(self):
        gender = secrets.choice(['male', 'female'])
        if gender == 'male':
            name = secrets.choice(self.russian_male_names)
            surname = secrets.choice(self.russian_surnames)
            patronymic = secrets.choice(self.russian_male_names) + 'ович'
        else:
            name = secrets.choice(self.russian_female_names)
            surname = secrets.choice(self.russian_surnames) + 'а'
            patronymic = secrets.choice(self.russian_male_names) + 'овна'
        return f"{surname} {name} {patronymic}"

    def generate_address(self):
        city = secrets.choice(self.cities)
        street = secrets.choice(self.streets)
        house = secrets.randbelow(200) + 1
        apartment = secrets.randbelow(300) + 1
        return f"г. {city}, ул. {street}, д. {house}, кв. {apartment}"

    def generate_inn(self):
        return ''.join(secrets.choice(string.digits) for _ in range(12))

    def generate_passport(self):
        series = ''.join(secrets.choice(string.digits) for _ in range(4))
        number = ''.join(secrets.choice(string.digits) for _ in range(6))
        return f"{series} №{number}"

    def generate_phone(self):
        code = secrets.choice(['495', '499', '812'])
        part1 = ''.join(secrets.choice(string.digits) for _ in range(3))
        part2 = ''.join(secrets.choice(string.digits) for _ in range(2))
        part3 = ''.join(secrets.choice(string.digits) for _ in range(2))
        return f"+7 ({code}) {part1}-{part2}-{part3}"

    def generate_date(self):
        year = secrets.choice(range(1950, 2005))
        month = secrets.choice(range(1, 13))
        day = secrets.choice(range(1, 29))
        return f"{day:02d}.{month:02d}.{year}"

    def generate_case_number(self):
        return f"2-{secrets.randbelow(10000)}/{secrets.choice(range(2020, 2026))}"

class EncryptionManager:
    def __init__(self):
        self.key_file = "encryption.key"
        self.key = self.load_or_generate_key()
        self.cipher = Fernet(self.key)
    
    def load_or_generate_key(self):
        if os.path.exists(self.key_file):
            with open(self.key_file, 'rb') as f:
                return f.read()
        else:
            key = Fernet.generate_key()
            with open(self.key_file, 'wb') as f:
                f.write(key)
            return key
    
    def encrypt(self, data):
        return self.cipher.encrypt(data.encode()).hex()
    
    def decrypt(self, encrypted_hex):
        return self.cipher.decrypt(bytes.fromhex(encrypted_hex)).decode()

class DatabaseManager:
    def __init__(self, encryption_manager=None):
        self.encryption_manager = encryption_manager
        self.conn = sqlite3.connect(DATABASE_FILE)
        self.init_database()
    
    def init_database(self):
        cursor = self.conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                original_file TEXT,
                processed_file TEXT,
                timestamp DATETIME
            )
        ''')
        
        if self.encryption_manager:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS encrypted_data (
                    id TEXT PRIMARY KEY,
                    document_id TEXT,
                    encrypted_text TEXT,
                    entity_type TEXT,
                    position_start INTEGER,
                    position_end INTEGER,
                    FOREIGN KEY (document_id) REFERENCES documents(id)
                )
            ''')
        else:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS replacements (
                    id TEXT PRIMARY KEY,
                    document_id TEXT,
                    original_text TEXT,
                    replacement_text TEXT,
                    entity_type TEXT,
                    position_start INTEGER,
                    position_end INTEGER,
                    FOREIGN KEY (document_id) REFERENCES documents(id)
                )
            ''')
        
        self.conn.commit()
    
    def save_document(self, doc_id, original_file, processed_file):
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT OR REPLACE INTO documents (id, original_file, processed_file, timestamp)
            VALUES (?, ?, ?, ?)
        ''', (doc_id, original_file, processed_file, datetime.now().isoformat()))
        self.conn.commit()
    
    def save_encrypted_mapping(self, entity_id, doc_id, encrypted_text, entity_type, start, end):
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT INTO encrypted_data (id, document_id, encrypted_text, entity_type, position_start, position_end)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (entity_id, doc_id, encrypted_text, entity_type, start, end))
        self.conn.commit()
    
    def save_replacement(self, entity_id, doc_id, original, replacement, entity_type, start, end):
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT INTO replacements (id, document_id, original_text, replacement_text, entity_type, position_start, position_end)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (entity_id, doc_id, original, replacement, entity_type, start, end))
        self.conn.commit()
    
    def get_entity_by_id(self, entity_id):
        cursor = self.conn.cursor()
        cursor.execute('SELECT * FROM encrypted_data WHERE id = ?', (entity_id,))
        return cursor.fetchone()
    
    def close(self):
        self.conn.close()

class Anonymizer:
    def __init__(self, mode='replace'):
        self.mode = mode
        
        if os.path.exists(MODEL_PATH):
            self.tokenizer = AutoTokenizer.from_pretrained(MODEL_PATH)
            self.model = AutoModelForTokenClassification.from_pretrained(MODEL_PATH)
            self.model.eval()
            self.id2label = self.model.config.id2label
            self.use_nlp_model = True
        else:
            print(f"Внимание: Модель не найдена в {MODEL_PATH}. Используется regex-режим.")
            self.use_nlp_model = False
        
        if mode == 'encrypt':
            self.encryption_manager = EncryptionManager()
            self.db_manager = DatabaseManager(self.encryption_manager)
        else:
            self.db_manager = DatabaseManager()
        
        self.data_generator = DataGenerator()
        self.stats = {'found': 0, 'processed': 0}
    
    def find_entities_by_patterns(self, text):
        entities = []
        
        patterns = {
            'PER': [
                r'\b[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2}\b',
                r'\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.\b'
            ],
            'ADDR': [
                r'г\.\s*[А-ЯЁ][а-яё-]+,\s*(?:ул\.|просп\.|пер\.|б-р)\s*[А-ЯЁ][а-яё-]+,\s*д\.\s*\d+(?:\s*кв\.\s*\d+)?',
                r'\bг\.\s*[А-ЯЁ][а-яё-]+\b'
            ],
            'DATE': [
                r'\b\d{1,2}\.\d{1,2}\.\d{4}\b',
                r'\b\d{1,2}\s+(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}\b'
            ],
            'PASS': [
                r'\b(?:паспорт\s+)?(?:\d{4}\s+)?№?\s*\d{6}\b',
                r'\bсерия\s*\d{4}\s*№\s*\d{6}\b'
            ],
            'INN': [
                r'\bИНН\s*\d{10}(?:\d{2})?\b',
                r'\b\d{10}(?:\d{2})?\b'
            ],
            'PHONE': [
                r'\b(?:\+7|8)[\s\-\(]?\d{3}[\s\-\)]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}\b'
            ],
            'CASE': [
                r'\bДело\s*№?\s*[\w\-/]+\b',
                r'\b№\s*[\w\-/]+\b'
            ]
        }
        
        for entity_type, regex_list in patterns.items():
            for pattern in regex_list:
                for match in re.finditer(pattern, text, re.IGNORECASE):
                    start, end = match.span()
                    if start > 0 and text[start-1].isalnum():
                        continue
                    if end < len(text) and text[end].isalnum():
                        continue
                    
                    entity_text = match.group()
                    entities.append({
                        'text': entity_text,
                        'type': entity_type,
                        'start': start,
                        'end': end
                    })
        
        return entities
    
    def extract_entities(self, text):
        if self.use_nlp_model:
            return self._extract_with_nlp(text)
        else:
            return self.find_entities_by_patterns(text)
    
    def _extract_with_nlp(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
        
        with torch.no_grad():
            outputs = self.model(**inputs)
            predictions = torch.argmax(outputs.logits, dim=2)
        
        tokens = self.tokenizer.convert_ids_to_tokens(inputs['input_ids'][0])
        labels = [self.id2label[p.item()] for p in predictions[0]]
        
        entities = []
        current_entity = []
        current_type = None
        current_start = 0
        
        for i, (token, label) in enumerate(zip(tokens, labels)):
            if token in ['[CLS]', '[SEP]', '[PAD]']:
                continue
            
            if label.startswith('B-'):
                if current_entity:
                    entity_text = self._tokens_to_text(current_entity)
                    if self._is_valid_entity(entity_text, current_type):
                        entities.append({
                            'text': entity_text,
                            'type': current_type,
                            'start': current_start,
                            'end': i
                        })
                
                current_entity = [token]
                current_type = label[2:]
                current_start = i
            
            elif label.startswith('I-') and current_type == label[2:]:
                current_entity.append(token)
            
            elif label == 'O':
                if current_entity:
                    entity_text = self._tokens_to_text(current_entity)
                    if self._is_valid_entity(entity_text, current_type):
                        entities.append({
                            'text': entity_text,
                            'type': current_type,
                            'start': current_start,
                            'end': i
                        })
                current_entity = []
                current_type = None
        
        if current_entity:
            entity_text = self._tokens_to_text(current_entity)
            if self._is_valid_entity(entity_text, current_type):
                entities.append({
                    'text': entity_text,
                    'type': current_type,
                    'start': current_start,
                    'end': len(tokens)
                })
        
        for entity in entities:
            entity['start_char'] = min(entity['start'] * 5, len(text))
            entity['end_char'] = min(entity['end'] * 5, len(text))
        
        return entities
    
    def _tokens_to_text(self, tokens):
        text = ""
        for token in tokens:
            if token.startswith('##'):
                text += token[2:]
            else:
                if text:
                    text += " "
                text += token
        return text.replace('▁', ' ').strip()
    
    def _is_valid_entity(self, text, entity_type):
        if not text or len(text) < 2:
            return False
        
        min_lengths = {
            'PER': 5,
            'ADDR': 10,
            'DATE': 4,
            'PASS': 8,
            'INN': 10,
            'PHONE': 10,
            'CASE': 3
        }
        
        min_len = min_lengths.get(entity_type, 3)
        if len(text) < min_len:
            return False
        
        if entity_type == 'PER':
            if ' ' not in text:
                return False
            words = text.split()
            if not all(word[0].isupper() for word in words if word):
                return False
        
        return True
    
    def process_document(self, input_path, output_path):
        doc_id = str(uuid.uuid4())[:8]
        doc = docx.Document(input_path)
        
        print(f"Обработка документа...")
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            if not text.strip():
                continue
            
            entities = self.extract_entities(text)
            self.stats['found'] += len(entities)
            
            if not entities:
                continue
            
            sorted_entities = sorted(entities, key=lambda x: x.get('start_char', text.find(x['text'])), reverse=True)
            
            for entity in sorted_entities:
                original = entity['text']
                etype = entity['type']
                
                if 'start_char' in entity:
                    pos = entity['start_char']
                    end_pos = entity['end_char']
                else:
                    pos = text.find(original)
                    end_pos = pos + len(original)
                
                if pos == -1:
                    continue
                
                if self.mode == 'replace':
                    replacement = self._generate_replacement(etype, original)
                    
                    entity_id = str(uuid.uuid4())[:8]
                    self.db_manager.save_replacement(
                        entity_id, doc_id, original, replacement, etype, pos, end_pos
                    )
                    
                    text = text[:pos] + replacement + text[end_pos:]
                    self.stats['processed'] += 1
                
                elif self.mode == 'encrypt':
                    encrypted = self.encryption_manager.encrypt(original)
                    entity_id = str(uuid.uuid4())[:8]
                    replacement = f"[ENC:{entity_id}]"
                    
                    self.db_manager.save_encrypted_mapping(
                        entity_id, doc_id, encrypted, etype, pos, end_pos
                    )
                    
                    text = text[:pos] + replacement + text[end_pos:]
                    self.stats['processed'] += 1
                
                elif self.mode == 'blur':
                    replacement = '█' * len(original)
                    text = text[:pos] + replacement + text[end_pos:]
                    self.stats['processed'] += 1
            
            if text != para.text:
                para.clear()
                para.add_run(text)
        
        self.db_manager.save_document(doc_id, os.path.basename(input_path), os.path.basename(output_path))
        
        doc.save(output_path)
        self.db_manager.close()
        
        return doc_id
    
    def _generate_replacement(self, entity_type, original=""):
        generators = {
            'PER': self.data_generator.generate_fio,
            'ADDR': self.data_generator.generate_address,
            'DATE': self.data_generator.generate_date,
            'PASS': self.data_generator.generate_passport,
            'INN': self.data_generator.generate_inn,
            'PHONE': self.data_generator.generate_phone,
            'CASE': self.data_generator.generate_case_number
        }
        
        generator = generators.get(entity_type)
        if generator:
            return generator()
        else:
            return f"[{entity_type}]"
    
    def decrypt_document(self, encrypted_file_path, output_path):
        if self.mode != 'encrypt' or not hasattr(self, 'encryption_manager'):
            print("Режим шифрования не активен")
            return False
        
        doc = docx.Document(encrypted_file_path)
        
        for para in doc.paragraphs:
            text = para.text
            
            enc_pattern = r'\[ENC:([a-f0-9]+)\]'
            for match in re.finditer(enc_pattern, text, re.IGNORECASE):
                entity_id = match.group(1)
                
                entity_data = self.db_manager.get_entity_by_id(entity_id)
                if entity_data:
                    encrypted_text = entity_data[2]
                    try:
                        decrypted = self.encryption_manager.decrypt(encrypted_text)
                        text = text.replace(match.group(0), decrypted)
                    except:
                        print(f"Ошибка расшифровки для ID: {entity_id}")
            
            if text != para.text:
                para.clear()
                para.add_run(text)
        
        doc.save(output_path)
        return True

def show_files():
    files = [f for f in os.listdir(INPUT_FOLDER) if f.lower().endswith('.docx')]
    
    if not files:
        doc = docx.Document()
        doc.add_paragraph("СУДЕБНОЕ РЕШЕНИЕ")
        doc.add_paragraph("Дело № 2-1234/2024")
        doc.add_paragraph("Истец: Иванов Иван Иванович")
        doc.add_paragraph("Дата рождения: 15.04.1985")
        doc.add_paragraph("Адрес: г. Москва, ул. Ленина, д. 10, кв. 5")
        doc.add_paragraph("Телефон: +7 (999) 123-45-67")
        doc.add_paragraph("ИНН: 123456789012")
        doc.add_paragraph("Паспорт: 1234 № 567890")
        doc.add_paragraph("Ответчик: Петрова Анна Сергеевна")
        
        filepath = os.path.join(INPUT_FOLDER, "пример.docx")
        doc.save(filepath)
        files = ["пример.docx"]
    
    return files

def main():
    os.makedirs(INPUT_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    
    print("=" * 60)
    print("СИСТЕМА ОБЕЗЛИЧИВАНИЯ ПЕРСОНАЛЬНЫХ ДАННЫХ")
    print("=" * 60)
    
    files = show_files()
    
    print("\nДоступные файлы:")
    for i, f in enumerate(files, 1):
        print(f"{i}. {f}")
    
    try:
        choice = int(input("\nВыберите файл: ")) - 1
        if choice < 0 or choice >= len(files):
            print("Неверный выбор")
            return
        
        print("\nРежимы обработки:")
        print("1. Замена на фиктивные данные")
        print("2. Шифрование с сохранением ключа")
        print("3. Размытие (блокировка)")
        
        mode_choice = input("Выберите режим (1-3): ")
        
        if mode_choice == '1':
            mode = 'replace'
            mode_name = "замена на фиктивные данные"
        elif mode_choice == '2':
            mode = 'encrypt'
            mode_name = "шифрование"
        elif mode_choice == '3':
            mode = 'blur'
            mode_name = "размытие"
        else:
            print("Неверный выбор")
            return
        
        input_file = files[choice]
        input_path = os.path.join(INPUT_FOLDER, input_file)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"обработанный_{input_file.replace('.docx', '')}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_name)
        
        print(f"\nРежим: {mode_name}")
        print("Обработка...")
        
        anonymizer = Anonymizer(mode=mode)
        doc_id = anonymizer.process_document(input_path, output_path)
        
        print(f"\n✓ Обработка завершена!")
        print(f"  Найдено сущностей: {anonymizer.stats['found']}")
        print(f"  Обработано сущностей: {anonymizer.stats['processed']}")
        print(f"  Сохранен файл: {output_path}")
        
        if mode == 'encrypt':
            print(f"  ID документа: {doc_id}")
            print(f"  Ключ шифрования: encryption.key")
            print(f"  База данных: {DATABASE_FILE}")
        
        if mode == 'encrypt':
            decrypt_choice = input("\nПротестировать расшифровку? (да/нет): ")
            if decrypt_choice.lower() in ['да', 'yes', 'y', 'д']:
                decrypt_path = os.path.join(OUTPUT_FOLDER, f"расшифрованный_{timestamp}.docx")
                if anonymizer.decrypt_document(output_path, decrypt_path):
                    print(f"  ✓ Расшифрованный файл: {decrypt_path}")
                else:
                    print("  ✗ Ошибка расшифровки")
        
    except Exception as e:
        print(f"\n✗ Ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
